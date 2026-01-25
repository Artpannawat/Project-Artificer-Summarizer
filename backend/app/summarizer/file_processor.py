import io
import os
from typing import Union
from fastapi import UploadFile, HTTPException

# พยายามนำเข้า Dependencies ที่เป็นตัวเลือก

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False



try:
    import docx2txt
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False


class FileProcessor:
    """
    คลาสสำหรับการประมวลผลไฟล์เอกสารรูปแบบต่างๆ (PDF, DOC, DOCX, TXT)
    และดึงเนื้อหาข้อความเพื่อการสรุป
    """
    
    SUPPORTED_FORMATS = {
        'application/pdf': 'pdf',
        'application/msword': 'doc',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'text/plain': 'txt'
    }
    
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    def __init__(self):
        pass
    
    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        ดึงเนื้อหาข้อความจากไฟล์ที่อัปโหลดโดยอิงตามรูปแบบไฟล์
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            str: Extracted text content
            
        Raises:
            HTTPException: If file format is not supported or extraction fails
        """
        # ตรวจสอบขนาดไฟล์
        content = await file.read()
        if len(content) > self.MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="ขนาดไฟล์เกิน 10MB")
        
        # รีเซ็ตตัวชี้ตำแหน่งไฟล์
        await file.seek(0)
        
        # ระบุรูปแบบไฟล์
        file_format = self._get_file_format(file)
        
        try:
            if file_format == 'pdf':
                text = await self._extract_from_pdf(content)
            elif file_format == 'docx':
                text = await self._extract_from_docx(content)
            elif file_format == 'doc':
                text = await self._extract_from_doc(content)
            elif file_format == 'txt':
                text = await self._extract_from_txt(content)
            else:
                raise HTTPException(status_code=400, detail="รองรับเฉพาะไฟล์ PDF, DOC, DOCX, TXT เท่านั้น")
            
            if not text or len(text.strip()) < 10:
                # Instead of error, return empty string to trigger OCR fallback in main.py
                # แทนที่จะแจ้ง Error ให้คืนค่าว่างเพื่อไปทำ OCR ต่อ
                return ""
            
            return text.strip()
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"เกิดข้อผิดพลาดในการประมวลผลไฟล์: {str(e)}")
    
    def _get_file_format(self, file: UploadFile) -> str:
        """
        ระบุรูปแบบไฟล์จาก Content Type หรือนามสกุลไฟล์
        """
        # ตรวจสอบ Content Type ก่อน
        if file.content_type in self.SUPPORTED_FORMATS:
            return self.SUPPORTED_FORMATS[file.content_type]
        
        # Fallback ไปที่นามสกุลไฟล์
        if file.filename:
            filename_lower = file.filename.lower()
            if filename_lower.endswith('.pdf'):
                return 'pdf'
            elif filename_lower.endswith('.docx'):
                return 'docx'
            elif filename_lower.endswith('.doc'):
                return 'doc'
            elif filename_lower.endswith('.txt'):
                return 'txt'
        
        raise HTTPException(status_code=400, detail="ไม่สามารถระบุประเภทไฟล์ได้")
    
    async def _extract_from_pdf(self, content: bytes) -> str:
        """ดึงข้อความจากไฟล์ PDF โดยใช้ pypdf (Lightweight for Vercel)"""
        try:
            import pypdf
        except ImportError:
             raise HTTPException(
                status_code=500, 
                detail="ไม่สามารถอ่านไฟล์ PDF ได้ เนื่องจากไม่มี pypdf installed"
            )
        
        try:
            text = ""
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                return text
                
            # Return empty string to signal scanned PDF
            return ""
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"ไม่สามารถอ่านไฟล์ PDF ได้: {str(e)}")
    
    async def _extract_from_docx(self, content: bytes) -> str:
        """ดึงข้อความจากไฟล์ DOCX"""
        if not HAS_DOCX2TXT and not HAS_PYTHON_DOCX:
            raise HTTPException(
                status_code=500, 
                detail="ไม่สามารถอ่านไฟล์ DOCX ได้ เนื่องจากไม่มี DOCX processing libraries"
            )
        
        try:
            text = ""
            
            # วิธีที่ 1: ใช้ docx2txt (ง่ายกว่าและเชื่อถือได้มากกว่า)
            if HAS_DOCX2TXT:
                try:
                    text = docx2txt.process(io.BytesIO(content))
                    if text.strip():
                        return text
                except Exception:
                    pass  # ลองใช้ python-docx เป็น fallback
            
            # วิธีที่ 2: ใช้ python-docx เป็น fallback
            if HAS_PYTHON_DOCX:
                try:
                    doc = Document(io.BytesIO(content))
                    paragraphs = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            paragraphs.append(paragraph.text.strip())
                    
                    # ดึงข้อความจากตารางด้วย
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    paragraphs.append(cell.text.strip())
                    
                    text = "\n".join(paragraphs)
                    if text.strip():
                        return text
                except Exception:
                    pass
            
            # ถ้าทั้งสองวิธีล้มเหลว
            raise HTTPException(
                status_code=500, 
                detail="ไม่สามารถอ่านเนื้อหาจากไฟล์ DOCX ได้ ไฟล์อาจเสียหายหรือไม่ถูกต้อง"
            )
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"ไม่สามารถอ่านไฟล์ DOCX ได้: {str(e)}")
    
    async def _extract_from_doc(self, content: bytes) -> str:
        """ดึงข้อความจากไฟล์ DOC (รูปแบบเก่า)"""
        if not HAS_DOCX2TXT:
            raise HTTPException(
                status_code=400, 
                detail="ไม่สามารถอ่านไฟล์ .doc ได้ กรุณาแปลงเป็น .docx หรือ PDF"
            )
        
        try:
            # สำหรับไฟล์ .doc เราจะลองใช้ docx2txt ซึ่งบางครั้งก็ใช้ได้
            # หมายเหตุ: การรองรับ .doc มีจำกัด แนะนำให้ผู้ใช้แปลงเป็น .docx
            text = docx2txt.process(io.BytesIO(content))
            
            if not text.strip():
                raise HTTPException(
                    status_code=400, 
                    detail="ไม่สามารถอ่านไฟล์ .doc ได้ กรุณาแปลงเป็น .docx หรือ PDF"
                )
            
            return text
            
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"ไฟล์ .doc อาจเสียหาย กรุณาแปลงเป็น .docx หรือ PDF: {str(e)}"
            )
    
    async def _extract_from_txt(self, content: bytes) -> str:
        """ดึงข้อความจากไฟล์ TXT"""
        try:
            # ลองใช้การเข้ารหัสแบบต่างๆ
            encodings = ['utf-8', 'utf-8-sig', 'cp874', 'iso-8859-1', 'windows-1252']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            # ถ้าการเข้ารหัสทั้งหมดล้มเหลว ให้ใช้ utf-8 พร้อมการจัดการข้อผิดพลาด
            text = content.decode('utf-8', errors='replace')
            return text
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"ไม่สามารถอ่านไฟล์ TXT ได้: {str(e)}")
    
    def validate_file(self, file: UploadFile) -> bool:
        """
        ตรวจสอบรูปแบบและขนาดไฟล์ที่อัปโหลด
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            bool: True if file is valid
            
        Raises:
            HTTPException: If file is invalid
        """
        # ตรวจสอบว่ามีไฟล์อยู่หรือไม่
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="ไม่พบไฟล์")
        
        # ตรวจสอบรูปแบบไฟล์
        try:
            self._get_file_format(file)
        except HTTPException:
            raise HTTPException(status_code=400, detail="รองรับเฉพาะไฟล์ PDF, DOC, DOCX, TXT เท่านั้น")
        
        return True